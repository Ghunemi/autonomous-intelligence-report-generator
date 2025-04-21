import os, textwrap
from typing import List, Dict
from datetime import date
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class ReportWriter:
    """
    Single responsibility: assemble a DOCX report from insights & sources.
    """
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        Path(self.output_dir).mkdir(exist_ok=True)

    def write(self, query: str, insights: Dict, sources: List[Dict]) -> str:
        filename = f"{query[:40].replace(' ', '_')}.docx"
        path = Path(self.output_dir) / filename

        doc = Document()
        doc.add_heading(query, level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(
            f"Date: {date.today().isoformat()}",
            style="Subtitle"
        ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

        # Trends
        doc.add_heading("Key Market Trends", level=1)
        for trend in insights["trends"]:
            doc.add_paragraph(trend, style="List Bullet")

        # Competitor table
        doc.add_heading("Competitor Landscape", level=1)
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Rank", "Company", "Note"
        for i, comp in enumerate(insights["competitors"], start=1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = comp["name"]
            row[2].text = textwrap.fill(comp["note"], 60)

        # Recommendations
        doc.add_heading("Strategic Recommendations", level=1)
        for rec in insights["recommendations"]:
            doc.add_paragraph(rec, style="List Number")

        # Sources
        doc.add_heading("Sources", level=1)
        for src in sources:
            doc.add_paragraph(src["url"], style="List Bullet")

        doc.save(path)
        print(f"[report] saved to {path}")
        return str(path)
